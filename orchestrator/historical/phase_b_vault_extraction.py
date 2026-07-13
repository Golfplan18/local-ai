"""Phase B — atomic note extraction from arbitrary vault documents.

Adapter over Phase 5 primitives (api_client, ChromaDB dedup, atomic
note writer) for extracting atomics from user-authored documents:
vault markdown, docx, doc, rtf, plain text.

Output: flat `~/Documents/vault/Engrams/` (no year subfolders).
Dedup: ChromaDB `atomic_dedup` collection (shared with Phase 5 corpus).
Manifest: `~/ora/data/phase-b-manifest.json` for resumability.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import subprocess
import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from orchestrator import runtime_paths as _rp
from orchestrator.historical.api_client import AnthropicClient
from orchestrator.historical.phase5_atomic_extraction import (
    _slugify,
    _yaml_escape,
    _strip_json_fences,
    _VALID_TYPES,
    DEFAULT_CHROMADB_PATH,
    DEFAULT_DEDUP_COLLECTION,
    DEDUP_SIM_THRESHOLD,
    EXTRACTION_MODEL,
    DedupResult,
    _open_dedup_collection,
)


# ---------------------------------------------------------------------------
# Defaults
# ---------------------------------------------------------------------------

DEFAULT_VAULT_ROOT_FLAT = str(_rp.vault_dir() / "Engrams")
DEFAULT_MANIFEST_PATH = str(_rp.DATA_DIR / "phase-b-manifest.json")
DEFAULT_REPORT_PATH = str(_rp.DATA_DIR / "phase-b-report.json")

# Extraction sizing for vault docs. Chapters and white papers are larger
# than chat pairs; we send up to ~24K chars per call (~6K tokens) and
# allow up to 40 atomics per call.
MAX_DOC_CHARS = 24_000
MAX_TOKENS_PER_CALL = 4096
MIN_DOC_CHARS = 200  # skip stubs

# Chunking for very long documents — split on h1/h2 boundaries when
# content exceeds MAX_DOC_CHARS, fall back to char-window split.
CHUNK_TARGET_CHARS = 20_000
CHUNK_OVERLAP_CHARS = 1_000


# ---------------------------------------------------------------------------
# Document loaders
# ---------------------------------------------------------------------------


def load_markdown(path: Path) -> str:
    """Load a markdown file, stripping YAML frontmatter."""
    text = path.read_text(encoding="utf-8", errors="replace")
    # Strip YAML frontmatter
    if text.startswith("---\n"):
        end = text.find("\n---\n", 4)
        if end > 0:
            text = text[end + 5 :]
    return text.strip()


def load_docx(path: Path) -> str:
    """Convert .docx to plain text via python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed")
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        s = p.text.strip()
        if s:
            # Heuristic markdown for headings
            style = (p.style.name or "").lower()
            if "heading 1" in style:
                parts.append(f"# {s}")
            elif "heading 2" in style:
                parts.append(f"## {s}")
            elif "heading 3" in style:
                parts.append(f"### {s}")
            else:
                parts.append(s)
    # Tables
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _textutil_convert(path: Path) -> str:
    """Run macOS textutil to convert a legacy document to plain text.

    textutil is a macOS system binary (ships with the OS) and is not available
    on Windows. Callers on Windows must install a cross-platform alternative
    for legacy binary `.doc` text extraction (e.g. `antiword` via `brew install
    antiword` on macOS / package managers on Linux / MSYS2 or Cygwin on Windows;
    `wv` for fuller extraction; or headless `LibreOffice` via
    `soffice --headless --convert-to txt`) or skip the file. NOTE: Mammoth is
    NOT a valid choice — it supports `.docx` only, not binary `.doc` (per the
    official `python-mammoth` documentation). NOTE: `.rtf` conversion is
    already supported through the installed `striprtf` dependency (see
    `load_rtf`); only legacy `.doc` is in D-09 scope. The guard fails loudly
    with a portable error rather than silently producing wrong behavior.
    """
    if os.name == "nt":
        raise RuntimeError(
            f"textutil is macOS-only and cannot convert {path} on Windows. "
            "Install a cross-platform legacy-`.doc` text extractor (e.g. "
            "`antiword`, `wv`, or headless `LibreOffice` via "
            "`soffice --headless --convert-to txt`) and add a Windows branch "
            "to load_doc. Note: `mammoth` is NOT a valid choice — it supports "
            "`.docx` only, not binary `.doc`. Note: `.rtf` is already "
            "supported through the installed `striprtf` dependency (load_rtf)."
        )
    out = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        capture_output=True, timeout=120,
    )
    if out.returncode != 0:
        raise RuntimeError(f"textutil failed: {out.stderr.decode(errors='replace')[:200]}")
    return out.stdout.decode("utf-8", errors="replace").strip()


def load_doc(path: Path) -> str:
    """Convert legacy .doc to text via macOS textutil."""
    return _textutil_convert(path)


def load_rtf(path: Path) -> str:
    """Convert .rtf to plain text via striprtf or textutil."""
    try:
        from striprtf.striprtf import rtf_to_text
        return rtf_to_text(path.read_text(encoding="utf-8", errors="replace")).strip()
    except ImportError:
        return _textutil_convert(path)


def load_document(path: Path) -> tuple[str, str]:
    """Return (text, source_format) for a document of any supported type."""
    suffix = path.suffix.lower()
    if suffix in (".md", ".markdown", ".txt"):
        return load_markdown(path) if suffix != ".txt" else path.read_text(encoding="utf-8", errors="replace").strip(), suffix.lstrip(".")
    if suffix == ".docx":
        return load_docx(path), "docx"
    if suffix == ".doc":
        return load_doc(path), "doc"
    if suffix == ".rtf":
        return load_rtf(path), "rtf"
    raise ValueError(f"unsupported format: {suffix}")


def chunk_text(text: str) -> list[tuple[int, str]]:
    """Split text into (chunk_idx, chunk_text) tuples ≤ CHUNK_TARGET_CHARS."""
    if len(text) <= CHUNK_TARGET_CHARS:
        return [(0, text)]
    # Try splitting on h1/h2 boundaries
    chunks: list[tuple[int, str]] = []
    pieces = re.split(r"(?m)^(?=##? )", text)
    if len(pieces) > 1 and all(len(p) <= CHUNK_TARGET_CHARS for p in pieces):
        return list(enumerate(p.strip() for p in pieces if p.strip()))
    # Fall back to char-window split with overlap
    idx = 0
    cursor = 0
    while cursor < len(text):
        end = min(cursor + CHUNK_TARGET_CHARS, len(text))
        # Try to break on a paragraph boundary if possible
        if end < len(text):
            last_blank = text.rfind("\n\n", cursor, end)
            if last_blank > cursor + CHUNK_TARGET_CHARS // 2:
                end = last_blank
        chunks.append((idx, text[cursor:end].strip()))
        idx += 1
        cursor = end - CHUNK_OVERLAP_CHARS if end < len(text) else end
    return chunks


# ---------------------------------------------------------------------------
# Sonnet extraction (single-document variant)
# ---------------------------------------------------------------------------


_SYSTEM_PROMPT_DOC = """\
You extract ATOMIC KNOWLEDGE NOTES from a single user-authored document \
(markdown, book chapter, white paper, mind map, etc.). Each atomic note \
states ONE clear, transferable claim — a complete declarative sentence, \
not a topic label.

Be SELECTIVE — only mint notes for genuinely insightful, transferable \
claims that someone might want to retrieve later. Skip trivial or \
well-known facts. A short note might yield 0-3 atomics; a substantive \
chapter or white paper might yield 10-30. Cap at 40 per call.

Categories (USER-SIDE — this is user-authored content):
- fact: verifiable empirical claim
- principle: generalizable rule about how something works
- definition: precise definition of a concept
- causal: assertion that one thing causes / prevents / enables another
- analogy: structural comparison between two domains
- evaluative: judgment with explicit criteria

Output format: JSON array. Reply with ONLY the JSON, no preamble or \
fences. If nothing worth minting, reply with `[]`.

Each object must have:
  - "title": ONE complete declarative sentence stating the claim
  - "type": one of the six category names above
  - "body": 1-3 proposition bullets explaining the claim, each with \
named actors (no "it" / "they" / passive voice — name what does what)
  - "confidence": "high" | "medium" | "low" (your confidence the claim \
is well-supported by what's actually in the text)

Example output:
[
  {
    "title": "Premature abstraction in code increases maintenance cost \
by forcing future changes through an interface that encoded the wrong \
assumptions",
    "type": "causal",
    "body": "- Abstractions designed before requirements stabilize lock \
in incorrect assumptions about variation\\n- Future changes must route \
through the abstraction's interface even when the underlying assumption \
is now wrong\\n- The cost compounds because each new requirement either \
fights the abstraction or accumulates workaround complexity",
    "confidence": "high"
  }
]"""


def call_sonnet_extract_document(
    document_text: str,
    *,
    client: AnthropicClient,
) -> tuple[list[dict], int, int, float, str]:
    text = document_text[:MAX_DOC_CHARS]
    body = (
        f"DOCUMENT:\n<<<\n{text}\n>>>\n\n"
        f"Extract atomic notes (JSON array, max 40 entries):"
    )
    result = client.call(
        system=_SYSTEM_PROMPT_DOC,
        user=body,
        model=EXTRACTION_MODEL,
        max_tokens=MAX_TOKENS_PER_CALL,
        temperature=0.0,
    )
    if result.error:
        return [], result.input_tokens, result.output_tokens, result.cost_usd, result.error
    raw = _strip_json_fences(result.text)
    try:
        parsed = json.loads(raw)
        if not isinstance(parsed, list):
            return [], result.input_tokens, result.output_tokens, result.cost_usd, "not a list"
    except json.JSONDecodeError as e:
        return [], result.input_tokens, result.output_tokens, result.cost_usd, f"json: {e}"
    cleaned: list[dict] = []
    for c in parsed:
        if not isinstance(c, dict):
            continue
        if c.get("type") not in _VALID_TYPES:
            continue
        if not c.get("title") or not c.get("body"):
            continue
        cleaned.append(c)
    return cleaned, result.input_tokens, result.output_tokens, result.cost_usd, ""


# ---------------------------------------------------------------------------
# Atomic candidate (document-sourced)
# ---------------------------------------------------------------------------


@dataclass
class DocAtomicCandidate:
    title: str
    note_type: str
    body: str
    confidence: str
    # Provenance
    source_path: str          # absolute path of the source file
    source_format: str        # "md" / "docx" / "doc" / "rtf" / "txt"
    source_label: str         # human-readable label (e.g., "We Too Book 1")
    chunk_idx: int            # 0 if doc fit in one call
    when: datetime            # source mtime or today


def build_doc_atomic_note(c: DocAtomicCandidate) -> str:
    today = datetime.now().strftime("%Y-%m-%d")
    when_str = c.when.strftime("%Y-%m-%d")
    rel_path = c.source_path.replace(os.path.expanduser("~/"), "~/")
    yaml_lines = [
        "---",
        "nexus:",
        "type: engram",
        "tags:",
        "  - atomic",
        f"  - {c.note_type}",
        f"date created: {when_str}",
        f"date modified: {today}",
        f"source_path: {_yaml_escape(rel_path)}",
        f"source_format: {c.source_format}",
        f"source_label: {_yaml_escape(c.source_label)}",
        f"source_side: user",
        f"extraction_model: {EXTRACTION_MODEL}",
        f"confidence: {c.confidence}",
        f"processed_at: {datetime.now().strftime('%Y-%m-%dT%H:%M:%S')}",
        "seen_count: 1",
        "---",
    ]
    yaml = "\n".join(yaml_lines) + "\n\n"
    body_lines = [f"# {c.title}\n", c.body.strip(), "", "## Source", "",
                  f"From {c.source_label} ({c.source_format}). Path: `{rel_path}`."]
    return yaml + "\n".join(body_lines) + "\n"


def _doc_atomic_uid(c: DocAtomicCandidate) -> str:
    h = hashlib.sha256()
    h.update(c.source_path.encode("utf-8"))
    h.update(b"|")
    h.update(c.title.encode("utf-8"))
    return "atomic-" + h.hexdigest()[:14]


def _doc_vault_path_for(c: DocAtomicCandidate, vault_root: str) -> Path:
    slug = _slugify(c.title) or "untitled"
    base = f"{c.when.strftime('%Y-%m-%d')}_{slug}.md"
    return Path(vault_root) / base


def write_doc_atomic_note(c: DocAtomicCandidate, vault_root: str) -> str:
    path = _doc_vault_path_for(c, vault_root)
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        path = path.with_name(f"{path.stem}-{_doc_atomic_uid(c)[-6:]}{path.suffix}")
    path.write_text(build_doc_atomic_note(c), encoding="utf-8")
    return str(path)


def _doc_embedding_text(c: DocAtomicCandidate) -> str:
    """Title + body concatenation used for ChromaDB dedup embedding."""
    return f"{c.title}\n\n{c.body}".strip()


# ---------------------------------------------------------------------------
# Dedup adapter (mirrors phase5 check_and_register but for DocAtomicCandidate)
# ---------------------------------------------------------------------------


def doc_check_and_register(
    c: DocAtomicCandidate,
    *,
    collection,
    embedder,
) -> DedupResult:
    text = _doc_embedding_text(c)
    embedding = embedder(text)
    # Query top-1
    res = collection.query(query_embeddings=[embedding], n_results=1)
    if res["ids"] and res["ids"][0]:
        sim = 1.0 - res["distances"][0][0]
        if sim >= DEDUP_SIM_THRESHOLD:
            existing_id = res["ids"][0][0]
            existing_md = res["metadatas"][0][0]
            new_count = int(existing_md.get("seen_count", 1)) + 1
            new_md = dict(existing_md)
            new_md["seen_count"] = new_count
            collection.update(ids=[existing_id], metadatas=[new_md])
            return DedupResult(
                is_duplicate=True,
                matched_id=existing_id,
                matched_similarity=sim,
                matched_path=existing_md.get("vault_path", ""),
            )
    return DedupResult(is_duplicate=False)


def doc_register_new(
    c: DocAtomicCandidate,
    vault_path: str,
    *,
    collection,
    embedder,
) -> str:
    text = _doc_embedding_text(c)
    embedding = embedder(text)
    uid = _doc_atomic_uid(c)
    metadata = {
        "title": c.title,
        "note_type": c.note_type,
        "source_side": "user",
        "source_path": c.source_path,
        "source_format": c.source_format,
        "source_label": c.source_label,
        "when": c.when.isoformat(),
        "vault_path": vault_path,
        "seen_count": 1,
    }
    collection.add(ids=[uid], embeddings=[embedding], metadatas=[metadata], documents=[text])
    return uid


# ---------------------------------------------------------------------------
# Per-document orchestration
# ---------------------------------------------------------------------------


@dataclass
class DocResult:
    source_path: str
    candidates_total: int = 0
    candidates_minted: int = 0
    candidates_dedup: int = 0
    chunks_processed: int = 0
    input_tokens: int = 0
    output_tokens: int = 0
    cost_usd: float = 0.0
    errors: list[str] = field(default_factory=list)


def _embedder_configured(text: str) -> list[float]:
    """Embed with the active provider, model, and dimension configuration."""
    from orchestrator.embedding import embed_text

    return embed_text(text)


def process_one_document(
    source_path: str,
    *,
    vault_root: str,
    client: AnthropicClient,
    collection,
    embedder,
    source_label: Optional[str] = None,
) -> DocResult:
    res = DocResult(source_path=source_path)
    p = Path(source_path)
    try:
        text, source_format = load_document(p)
    except Exception as e:
        res.errors.append(f"load: {e}")
        return res
    if len(text) < MIN_DOC_CHARS:
        res.errors.append(f"skipped: only {len(text)} chars")
        return res

    label = source_label or p.stem
    when = datetime.fromtimestamp(p.stat().st_mtime)
    chunks = chunk_text(text)
    res.chunks_processed = len(chunks)

    for chunk_idx, chunk_text_str in chunks:
        candidates, in_t, out_t, cost, err = call_sonnet_extract_document(
            chunk_text_str, client=client
        )
        res.input_tokens += in_t
        res.output_tokens += out_t
        res.cost_usd += cost
        if err:
            res.errors.append(f"chunk {chunk_idx}: {err}")
            continue
        for raw in candidates:
            res.candidates_total += 1
            try:
                cand = DocAtomicCandidate(
                    title=raw["title"].strip(),
                    note_type=raw["type"],
                    body=raw["body"].strip(),
                    confidence=raw.get("confidence", "medium"),
                    source_path=str(p.absolute()),
                    source_format=source_format,
                    source_label=label,
                    chunk_idx=chunk_idx,
                    when=when,
                )
            except (KeyError, AttributeError) as e:
                res.errors.append(f"build: {e}")
                continue
            try:
                dedup = doc_check_and_register(cand, collection=collection, embedder=embedder)
            except Exception as e:
                res.errors.append(f"dedup: {e}")
                continue
            if dedup.is_duplicate:
                res.candidates_dedup += 1
                continue
            try:
                vp = write_doc_atomic_note(cand, vault_root)
                doc_register_new(cand, vp, collection=collection, embedder=embedder)
                res.candidates_minted += 1
            except Exception as e:
                res.errors.append(f"write: {e}")
    return res


# ---------------------------------------------------------------------------
# Manifest + orchestrator
# ---------------------------------------------------------------------------


def _load_manifest(path: str) -> dict:
    if os.path.exists(path):
        with open(path) as fh:
            return json.load(fh)
    return {
        "version": 1,
        "created_at": datetime.now().isoformat(),
        "completed_files": {},
        "totals": {
            "files_processed": 0,
            "candidates_total": 0,
            "candidates_minted": 0,
            "candidates_dedup": 0,
            "input_tokens": 0,
            "output_tokens": 0,
            "cost_usd": 0.0,
            "errors": 0,
        },
        "last_updated": datetime.now().isoformat(),
    }


def _save_manifest(manifest: dict, path: str) -> None:
    manifest["last_updated"] = datetime.now().isoformat()
    tmp = path + ".tmp"
    with open(tmp, "w") as fh:
        json.dump(manifest, fh, indent=2)
    os.replace(tmp, path)


def run_phase_b(
    file_paths: list[str],
    *,
    vault_root: str = DEFAULT_VAULT_ROOT_FLAT,
    chromadb_path: str = DEFAULT_CHROMADB_PATH,
    dedup_collection: str = DEFAULT_DEDUP_COLLECTION,
    manifest_path: str = DEFAULT_MANIFEST_PATH,
    max_workers: int = 8,
    label_provider=None,
) -> None:
    client = AnthropicClient()
    collection = _open_dedup_collection(chromadb_path, dedup_collection)
    manifest = _load_manifest(manifest_path)
    completed = manifest["completed_files"]
    todo = [p for p in file_paths if p not in completed]
    print(f"Phase B: {len(file_paths)} total, {len(todo)} pending, {len(completed)} resumed")
    sys.stdout.flush()
    if not todo:
        print("nothing to do")
        return

    t0 = time.time()
    processed = 0
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        futures = {
            ex.submit(
                process_one_document,
                p,
                vault_root=vault_root,
                client=client,
                collection=collection,
                embedder=_embedder_configured,
                source_label=(label_provider(p) if label_provider else None),
            ): p
            for p in todo
        }
        for fut in as_completed(futures):
            p = futures[fut]
            try:
                res = fut.result()
            except Exception as e:
                res = DocResult(source_path=p)
                res.errors.append(f"executor: {e}")
            completed[p] = {
                "candidates_total": res.candidates_total,
                "candidates_minted": res.candidates_minted,
                "candidates_dedup": res.candidates_dedup,
                "input_tokens": res.input_tokens,
                "output_tokens": res.output_tokens,
                "cost_usd": res.cost_usd,
                "chunks_processed": res.chunks_processed,
                "errors": res.errors,
            }
            tot = manifest["totals"]
            tot["files_processed"] += 1
            tot["candidates_total"] += res.candidates_total
            tot["candidates_minted"] += res.candidates_minted
            tot["candidates_dedup"] += res.candidates_dedup
            tot["input_tokens"] += res.input_tokens
            tot["output_tokens"] += res.output_tokens
            tot["cost_usd"] += res.cost_usd
            tot["errors"] += len(res.errors)
            processed += 1
            if processed % 5 == 0 or processed == len(todo):
                _save_manifest(manifest, manifest_path)
                elapsed = time.time() - t0
                rate = processed / elapsed if elapsed > 0 else 0.0
                remaining = (len(todo) - processed) / rate if rate > 0 else 0
                print(
                    f"  [{processed}/{len(todo)}] "
                    f"minted={tot['candidates_minted']} dedup={tot['candidates_dedup']} "
                    f"cost=${tot['cost_usd']:.2f} rate={rate:.2f}/s eta={int(remaining/60)}m"
                )
                sys.stdout.flush()
    _save_manifest(manifest, manifest_path)
    print(f"\nDONE — totals: {manifest['totals']}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def gather_files(
    roots: list[str],
    extensions: tuple[str, ...] = (".md", ".docx", ".doc", ".rtf", ".txt"),
    excludes: tuple[str, ...] = (),
) -> list[str]:
    """Walk roots, return absolute paths to files with allowed extensions,
    skipping any path containing an exclude segment."""
    found: list[str] = []
    for root in roots:
        rp = Path(root).expanduser()
        if not rp.exists():
            print(f"WARN: {rp} does not exist", file=sys.stderr)
            continue
        if rp.is_file():
            if rp.suffix.lower() in extensions:
                found.append(str(rp.absolute()))
            continue
        for path in rp.rglob("*"):
            if not path.is_file():
                continue
            if path.suffix.lower() not in extensions:
                continue
            s = str(path)
            if any(seg in s for seg in excludes):
                continue
            found.append(str(path.absolute()))
    return sorted(set(found))


def main(argv: Optional[list[str]] = None) -> int:
    p = argparse.ArgumentParser(description="Phase B vault-document atomic extraction")
    p.add_argument("--roots", nargs="+", required=True, help="Root dirs or files to scan")
    p.add_argument("--exclude", nargs="*", default=[], help="Path substrings to exclude")
    p.add_argument("--vault-root", default=DEFAULT_VAULT_ROOT_FLAT)
    p.add_argument("--manifest", default=DEFAULT_MANIFEST_PATH)
    p.add_argument("--max-workers", type=int, default=8)
    p.add_argument("--dry-run", action="store_true", help="List files only")
    args = p.parse_args(argv)

    files = gather_files(args.roots, excludes=tuple(args.exclude))
    print(f"Found {len(files)} files")
    if args.dry_run:
        for f in files[:50]:
            print(f"  {f}")
        if len(files) > 50:
            print(f"  ... ({len(files) - 50} more)")
        return 0

    run_phase_b(
        files,
        vault_root=args.vault_root,
        manifest_path=args.manifest,
        max_workers=args.max_workers,
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
